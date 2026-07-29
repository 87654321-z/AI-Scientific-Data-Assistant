"""数据导入页面。"""

import pandas as pd
import streamlit as st
from docx import Document


st.title("数据导入")
st.write("请选择实验记录图片、CSV、Excel 或 Word 文件。当前阶段只做上传和预览，不进行 OCR 识别。")
st.caption("建议单个文件不超过 20 MB。当前没有修改 Streamlit 的默认上传大小限制。")

uploaded_files = st.file_uploader(
    "选择文件",
    type=["png", "jpg", "jpeg", "csv", "xlsx", "docx"],
    accept_multiple_files=True,
)

if not uploaded_files:
    st.info("尚未选择文件。")

for uploaded_file in uploaded_files or []:
    file_extension = uploaded_file.name.rsplit(".", maxsplit=1)[-1].lower()
    file_size_mb = uploaded_file.size / 1024 / 1024

    st.divider()
    st.subheader(uploaded_file.name)
    st.write(f"文件类型：{uploaded_file.type or file_extension}")
    st.write(f"文件大小：{file_size_mb:.2f} MB")
    st.success("上传成功。")

    if file_extension in {"png", "jpg", "jpeg"}:
        st.image(uploaded_file, width=600)

    elif file_extension == "csv":
        try:
            csv_data = pd.read_csv(uploaded_file)
            st.write("CSV 前 20 行预览：")
            st.dataframe(csv_data.head(20), use_container_width=True)
        except Exception:
            st.error("无法读取这个 CSV 文件。请确认文件内容是有效的表格数据后再试。")

    elif file_extension == "xlsx":
        try:
            with pd.ExcelFile(uploaded_file, engine="openpyxl") as excel_file:
                sheet_names = excel_file.sheet_names
                first_sheet = sheet_names[0]
                excel_data = pd.read_excel(excel_file, sheet_name=first_sheet)

            st.write("工作表名称：" + "、".join(sheet_names))
            st.write(f"“{first_sheet}”前 20 行预览：")
            st.dataframe(excel_data.head(20), use_container_width=True)
        except Exception:
            st.error("无法读取这个 Excel 文件。请确认文件没有损坏，并且包含可读取的工作表。")

    elif file_extension == "docx":
        try:
            word_document = Document(uploaded_file)
            text_content = "\n".join(
                paragraph.text for paragraph in word_document.paragraphs if paragraph.text.strip()
            )

            st.write("Word 文字预览（前 1,000 个字符）：")
            st.text_area(
                "Word 文本内容",
                value=text_content[:1000] or "文档中没有可读取的文字内容。",
                height=250,
                disabled=True,
            )

            if word_document.tables:
                st.write(f"检测到 {len(word_document.tables)} 个表格：")
                for table_number, word_table in enumerate(word_document.tables, start=1):
                    table_rows = [
                        [cell.text for cell in row.cells] for row in word_table.rows
                    ]
                    st.write(f"表格 {table_number}：")
                    st.dataframe(pd.DataFrame(table_rows), use_container_width=True)
            else:
                st.write("该 Word 文档中没有表格。")
        except Exception:
            st.error("无法读取这个 Word 文件。请确认文件没有损坏，并且是 .docx 格式。")
